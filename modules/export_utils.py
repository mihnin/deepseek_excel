# modules/export_utils.py

import pandas as pd
import json
import csv
import io
import base64
from docx import Document
import streamlit as st

class ExportManager:
    @staticmethod
    def to_excel(df, filename="export.xlsx"):
        """Экспорт в Excel с форматированием"""
        output = io.BytesIO()
        
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            df.to_excel(writer, index=False, sheet_name='Результаты')
            
            # Получаем созданный лист
            workbook = writer.book
            worksheet = writer.sheets['Результаты']
            
            # Форматирование: автоподбор ширины столбцов
            for i, col in enumerate(df.columns):
                max_length = max(
                    df[col].astype(str).apply(len).max(),
                    len(str(col))
                ) + 2
                worksheet.column_dimensions[chr(65 + i)].width = min(max_length, 50)
        
        return output.getvalue()
    
    @staticmethod
    def to_csv(df, filename="export.csv", encoding="utf-8-sig"):
        """Экспорт в CSV с указанной кодировкой"""
        output = io.StringIO()
        df.to_csv(output, index=False, encoding=encoding)
        return output.getvalue().encode(encoding)
    
    @staticmethod
    def to_json(df, filename="export.json", orient="records"):
        """Экспорт в JSON с настраиваемым форматом"""
        json_data = df.to_json(orient=orient, force_ascii=False, indent=4)
        return json_data.encode("utf-8")
    
    @staticmethod
    def to_word(df, table_analysis=None, filename="export.docx"):
        """Экспорт в документ Word"""
        doc = Document()
        
        # Добавляем заголовок
        doc.add_heading('Отчет по анализу данных', level=1)
        
        # Если есть анализ всей таблицы, добавляем его
        if table_analysis:
            doc.add_heading('Общий анализ', level=2)
            doc.add_paragraph(table_analysis)
        
        # Добавляем таблицу с результатами
        doc.add_heading('Результаты построчного анализа', level=2)
        
        # Создаем таблицу в Word
        table = doc.add_table(rows=1, cols=len(df.columns))
        table.style = 'Table Grid'
        
        # Заполняем заголовки
        for i, column in enumerate(df.columns):
            table.cell(0, i).text = str(column)
        
        # Заполняем данные
        for _, row in df.iterrows():
            cells = table.add_row().cells
            for i, value in enumerate(row):
                cells[i].text = str(value)
        
        # Сохраняем в BytesIO
        output = io.BytesIO()
        doc.save(output)
        return output.getvalue()
    
    @staticmethod
    def create_download_button(data, filename, label, mime_type):
        """Создает кнопку для скачивания в интерфейсе Streamlit"""
        return st.download_button(
            label=label,
            data=data,
            file_name=filename,
            mime=mime_type
        )