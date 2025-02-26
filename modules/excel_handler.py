# modules/excel_handler.py
import pandas as pd
import numpy as np
from io import BytesIO
import io  # Добавляем импорт io для работы с потоками ввода-вывода
from typing import Dict, List, Tuple, Optional, Any, Union
import streamlit as st
from docx import Document

class ExcelHandler:
    """
    Класс для работы с Excel файлами - загрузка, обработка, анализ и сохранение.
    """
    
    @staticmethod
    def load_excel(file) -> pd.DataFrame:
        """
        Загружает Excel файл и возвращает DataFrame.
        
        Args:
            file: Файл Excel (BytesIO или путь)
            
        Returns:
            pd.DataFrame: Загруженные данные
            
        Raises:
            ValueError: Если не удалось прочитать файл
        """
        try:
            df = pd.read_excel(file)
            # Присваиваем имя для дальнейшего использования
            df.name = getattr(file, 'name', 'Unnamed Excel File')
            return df
        except Exception as e:
            raise ValueError(f"Ошибка при чтении Excel файла: {e}")
    
    @staticmethod
    def analyze_dataframe(df: pd.DataFrame) -> Dict[str, Any]:
        """
        Анализирует DataFrame и возвращает статистику.
        
        Args:
            df (pd.DataFrame): DataFrame для анализа
            
        Returns:
            Dict[str, Any]: Статистика по DataFrame
        """
        stats = {
            "shape": df.shape,
            "rows": df.shape[0],
            "columns": df.shape[1],
            "column_names": list(df.columns),
            "dtypes": {col: str(dtype) for col, dtype in df.dtypes.items()},
            "missing_values": {col: int(df[col].isna().sum()) for col in df.columns},
            "missing_percentage": {col: round(df[col].isna().sum() / len(df) * 100, 2) for col in df.columns},
            "has_issues": False,
            "issues": []
        }
        
        # Проверка на проблемы с данными
        for col in df.columns:
            # Проверка на большое количество пропусков
            if df[col].isna().sum() > df.shape[0] * 0.5:  # Более 50% пропусков
                stats["has_issues"] = True
                stats["issues"].append(f"Столбец '{col}' содержит более 50% пропущенных значений")
            
            # Проверка на однородность данных для строковых столбцов
            if df[col].dtype == 'object' and df[col].nunique() == 1 and df.shape[0] > 10:
                stats["has_issues"] = True
                stats["issues"].append(f"Столбец '{col}' содержит одинаковые значения во всех строках")
        
        return stats
    
    @staticmethod
    def clean_dataframe(df: pd.DataFrame, options: Dict[str, bool] = None) -> pd.DataFrame:
        """
        Очищает DataFrame от проблемных данных.
        
        Args:
            df (pd.DataFrame): DataFrame для очистки
            options (Dict[str, bool], optional): Опции очистки
                - remove_duplicates: Удалять дубликаты
                - fill_na: Заполнять пропущенные значения
                - normalize_text: Нормализовать текстовые данные
            
        Returns:
            pd.DataFrame: Очищенный DataFrame
        """
        if options is None:
            options = {
                "remove_duplicates": True,
                "fill_na": True,
                "normalize_text": True
            }
        
        # Создаем копию, чтобы не модифицировать оригинал
        cleaned_df = df.copy()
        
        # Удаление дубликатов
        if options.get("remove_duplicates", True):
            cleaned_df = cleaned_df.drop_duplicates()
        
        # Заполнение пропущенных значений
        if options.get("fill_na", True):
            # Для числовых столбцов используем медиану
            for col in cleaned_df.select_dtypes(include=['number']).columns:
                cleaned_df[col] = cleaned_df[col].fillna(cleaned_df[col].median())
            
            # Для текстовых столбцов используем пустую строку
            for col in cleaned_df.select_dtypes(include=['object']).columns:
                cleaned_df[col] = cleaned_df[col].fillna('')
            
            # Для дат используем самую частую дату или текущую
            for col in cleaned_df.select_dtypes(include=['datetime']).columns:
                if cleaned_df[col].notna().any():
                    most_common = cleaned_df[col].mode()[0]
                    cleaned_df[col] = cleaned_df[col].fillna(most_common)
                else:
                    cleaned_df[col] = cleaned_df[col].fillna(pd.Timestamp.now())
        
        # Нормализация текстовых данных
        if options.get("normalize_text", True):
            for col in cleaned_df.select_dtypes(include=['object']).columns:
                # Удаление лишних пробелов
                cleaned_df[col] = cleaned_df[col].astype(str).str.strip()
        
        return cleaned_df
    
    @staticmethod
    def suggest_target_columns(df: pd.DataFrame) -> List[str]:
        """
        Предлагает столбцы, которые могут быть целевыми для анализа.
        
        Args:
            df (pd.DataFrame): DataFrame для анализа
            
        Returns:
            List[str]: Список рекомендуемых целевых столбцов
        """
        suggestions = []
        
        # Анализ столбцов для выявления потенциальных целевых
        for col in df.columns:
            # Текстовые столбцы с достаточно длинными значениями могут содержать анализируемый текст
            if df[col].dtype == 'object':
                # Вычисляем среднюю длину текста в столбце
                avg_length = df[col].astype(str).str.len().mean()
                
                # Если средняя длина текста больше 50 символов, это может быть целевой столбец
                if avg_length > 50:
                    suggestions.append(col)
            
            # Столбцы с именами, указывающими на комментарии, отзывы и т.д.
            text_indicators = ["комментарий", "отзыв", "текст", "описание", "примечание", 
                              "comment", "review", "text", "description", "note"]
            
            if any(indicator in col.lower() for indicator in text_indicators):
                if col not in suggestions:
                    suggestions.append(col)
        
        return suggestions
    
    @staticmethod
    def suggest_additional_columns(df: pd.DataFrame, target_column: str) -> List[str]:
        """
        Предлагает столбцы, которые могут дать дополнительный контекст.
        
        Args:
            df (pd.DataFrame): DataFrame для анализа
            target_column (str): Выбранный целевой столбец
            
        Returns:
            List[str]: Список рекомендуемых дополнительных столбцов
        """
        suggestions = []
        
        # Столбцы с метаданными часто полезны как контекст
        metadata_indicators = ["дата", "время", "автор", "категория", "рейтинг", "оценка",
                              "date", "time", "author", "category", "rating", "score"]
        
        for col in df.columns:
            if col == target_column:
                continue
                
            # Столбцы с именами, указывающими на метаданные
            if any(indicator in col.lower() for indicator in metadata_indicators):
                suggestions.append(col)
            
            # Столбцы с небольшим количеством уникальных значений могут быть категориями
            if df[col].dtype == 'object' and df[col].nunique() < 10:
                if col not in suggestions:
                    suggestions.append(col)
            
            # Числовые столбцы с небольшим диапазоном тоже могут быть категориями или рейтингами
            if pd.api.types.is_numeric_dtype(df[col]) and df[col].nunique() < 10:
                if col not in suggestions:
                    suggestions.append(col)
        
        return suggestions
    
    @staticmethod
    def save_to_excel(df: pd.DataFrame, include_index: bool = False) -> BytesIO:
        """
        Сохраняет DataFrame в Excel-файл в памяти.
        
        Args:
            df (pd.DataFrame): DataFrame для сохранения
            include_index (bool): Включать индекс в файл
            
        Returns:
            BytesIO: Объект с Excel-файлом в памяти
        """
        output = BytesIO()
        
        try:
            with pd.ExcelWriter(output, engine="openpyxl") as writer:
                df.to_excel(writer, index=include_index, sheet_name="Results")
            
            # Перемещаем указатель в начало файла
            output.seek(0)
            return output
            
        except Exception as e:
            raise ValueError(f"Ошибка при сохранении в Excel: {e}")
    
    # Добавление методов из excel_utils.py
    @staticmethod
    def to_excel(df, filename="export.xlsx"):
        """Экспорт в Excel с форматированием"""
        output = io.BytesIO()
        
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            df.to_excel(writer, index=False, sheet_name='Результаты')
            
            # Получаем созданный лист
            workbook = writer.book
            worksheet = writer.sheets['Результаты']
            
            # Форматирование: автоподбор ширины столбцов
            for i, col in enumerate(df.columns):
                max_length = max(
                    df[col].astype(str).apply(len).max(),
                    len(str(col))
                ) + 2
                worksheet.column_dimensions[chr(65 + i)].width = min(max_length, 50)
        
        return output.getvalue()
    
    @staticmethod
    def to_csv(df, filename="export.csv", encoding="utf-8-sig"):
        """Экспорт в CSV с указанной кодировкой"""
        output = io.StringIO()
        df.to_csv(output, index=False, encoding=encoding)
        return output.getvalue().encode(encoding)
    
    @staticmethod
    def to_json(df, filename="export.json", orient="records"):
        """Экспорт в JSON с настраиваемым форматом"""
        json_data = df.to_json(orient=orient, force_ascii=False, indent=4)
        return json_data.encode("utf-8")
    
    @staticmethod
    def to_word(df, table_analysis=None, filename="export.docx"):
        """Экспорт в документ Word"""
        doc = Document()
        
        # Добавляем заголовок
        doc.add_heading('Отчет по анализу данных', level=1)
        
        # Если есть анализ всей таблицы, добавляем его
        if table_analysis:
            doc.add_heading('Общий анализ', level=2)
            doc.add_paragraph(table_analysis)
        
        # Добавляем таблицу с результатами
        doc.add_heading('Результаты построчного анализа', level=2)
        
        # Создаем таблицу в Word
        table = doc.add_table(rows=1, cols=len(df.columns))
        table.style = 'Table Grid'
        
        # Заполняем заголовки
        for i, column in enumerate(df.columns):
            table.cell(0, i).text = str(column)
        
        # Заполняем данные
        for _, row in df.iterrows():
            cells = table.add_row().cells
            for i, value in enumerate(row):
                cells[i].text = str(value)
        
        # Сохраняем в BytesIO
        output = io.BytesIO()
        doc.save(output)
        return output.getvalue()
    
    @staticmethod
    def create_download_button(data, filename, label, mime_type):
        """Создает кнопку для скачивания в интерфейсе Streamlit"""
        return st.download_button(
            label=label,
            data=data,
            file_name=filename,
            mime=mime_type
        )